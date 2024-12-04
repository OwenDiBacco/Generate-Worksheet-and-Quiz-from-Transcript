import os
import io
import json
import docx

import google.generativeai as genai
from dotenv import load_dotenv


from spire.doc import *
from spire.doc.common import *

import google.generativeai as genai
import webbrowser
from apiclient import discovery
from oauth2client.file import Storage
from httplib2 import Http
from oauth2client import tools, client

import tkinter as tk
from tkinter import filedialog
from PIL import Image, ImageTk

import google.generativeai as genai

import webbrowser
from apiclient import discovery
from oauth2client.file import Storage
from httplib2 import Http
from oauth2client import tools, client


class App:
    def __init__(self):
        self.root = tk.Tk()
        self.root.geometry('200x250')
        self.root.title("Txt File Selector")

        original_image = Image.open('images/txt_image.png')
        resized_image = original_image.resize((50, 50), Image.LANCZOS)
        zip_icon = ImageTk.PhotoImage(resized_image)

        zip_button = tk.Button(self.root, image=zip_icon, command=self.select_txt_file, borderwidth=0)
        zip_button.pack(pady=10)


        description_label = tk.Label(self.root, text="Open Txt File")
        description_label.pack()

        close_button = tk.Button(self.root, text="Run", command=self.close_app)
        close_button.pack(side=tk.BOTTOM, pady=10)

        self.root.mainloop()

    def select_txt_file(self):
        selected_file = filedialog.askopenfilename(
            title="Select a .txt file",
            filetypes=[("Text Files", "*.txt")]  # filter to only show .txt files
        )
        
        # return the selected file path
        if selected_file:
            self.txt_file = selected_file
            print(self.txt_file)

        else:
            print("no file selected.")
            self.txt_file = None
            print(self.txt_file)
    

    def close_app(self):
        self.root.destroy()


def read_txt_file(txt_path):
    with open(txt_path, 'r', encoding='utf-8') as file:
        return file.read()


def prompt_genai(prompt):
    text = ''
    load_dotenv() # Load variables from .env
    api_key = os.getenv("API_KEY")  # Access the API key
    genai.configure(api_key=api_key) # Replace with your own Gemini API Key
    model = genai.GenerativeModel("gemini-1.5-flash")
    responses = model.generate_content([prompt])
    for response in responses:
        text += ' ' + response.text
    
    return text


def select_file():
    app = App()
    txt_file_path = app.txt_file
    return txt_file_path


'''
JSON Data is for the Kahoot
'''
def generate_json_data(txt_file_path):
    txt_file_content = read_txt_file(txt_file_path)
    prompt = read_txt_file(".\\Kahoot-prompt.txt")
    form_prompt = f"{prompt} Notes: {txt_file_content}"
    form_question_response = prompt_genai(form_prompt)
    forms_questions_json = form_question_response.strip() # removes leading and trailing whitespace
    forms_questions_json = forms_questions_json[7:-3].strip() # removes the first 7 characters and the last 3 characters
    return forms_questions_json


'''
Code in function taken from Tanmay
'''
def generate_form(forms_questions_json):
    forms_questions_json = json.loads(forms_questions_json) # creates a json object

    SCOPES = [
        "https://www.googleapis.com/auth/forms.body",
        "https://www.googleapis.com/auth/forms.responses.readonly"
    ]

    DISCOVERY_DOC = "https://forms.googleapis.com/$discovery/rest?version=v1"

    store = Storage("token.json") # token is Tanmay's
    creds = store.get() # gets credentials from the token

    if not creds or creds.invalid:
        flow = client.flow_from_clientsecrets("client_secrets.json", SCOPES) # gets client secret; client secret is Tanmay's
        creds = tools.run_flow(flow, store)

    form_service = discovery.build( # creates a service object to interact with the Google Forms API
        "forms", "v1", http=creds.authorize(Http()), discoveryServiceUrl=DISCOVERY_DOC, static_discovery=False
    )

    form = { # initalizes the Google Form 
        "info": {
            "title": "Generated Quiz",
        }
    }

    result = form_service.forms().create(body=form).execute() # creates the Google Form
    update_requests = [
        {
            "updateSettings": {
                "settings": {"quizSettings": {"isQuiz": True}},
                "updateMask": "quizSettings.isQuiz",
            }
        }
    ]
    
    # creates the questions using JSON
    question_index = 0
    for question_data in forms_questions_json:
        question_text = question_data.get("question", "")
        options = [{"value": option} for option in question_data.get("options", [])]
        correct_answer = question_data.get("correctAnswer", "A")
        parsed_correct_answer = ord(correct_answer.lower()) - ord('a')
        question_request = {
            "createItem": {
                "item": {
                    "title": question_text,
                    "questionItem": {
                        "question": {
                            "required": True,
                            "grading": {
                                "pointValue": 1,
                                "correctAnswers": {
                                    "answers": [{"value": options[parsed_correct_answer]["value"]}]
                                },
                                "whenRight": {"text": "You got it!"},
                                "whenWrong": {"text": "Sorry, that's wrong"}
                            },
                            "choiceQuestion": {
                                "type": "RADIO",
                                "options": options,
                                "shuffle": True
                            }
                        }
                    }
                },
                "location": {
                    "index": question_index 
                }
            }
        }
        update_requests.append(question_request)
        question_index += 1
    
    # creates all the questions in the Google Form
    form_service.forms().batchUpdate(formId=result["formId"], body={"requests": update_requests}).execute()
    # Opens the Google Form
    form_url = f"https://docs.google.com/forms/d/{result['formId']}/edit"
    webbrowser.open(form_url)
    form_txt = io.BytesIO()
    form_txt.write(form_url.encode('utf-8'))
    form_txt.seek(0)
    generated_files['form'] = form_txt


def generate_questions(txt_file, number_of_questions = 10):    
    txt_content = read_txt_file(txt_file)
    prompt = read_txt_file(".\\Worksheet-prompt.txt")
    prompt += f'Transcript: {txt_content}\nNumber of questions: {number_of_questions}'
    response = prompt_genai(prompt)
    response = response.strip() # removes leading and trailing whitespace
    questions = response[7:-3].strip() # removes the first 7 characters and the last 3 characters
    return questions


def generate_worksheet(file_name, questions):
    def generate_table(section, document):
        table = Table(document, True)
        table.PreferredWidth = PreferredWidth(WidthType.Percentage, int(100))
        table.TableFormat.Borders.BorderType = BorderStyle.Single
        table.TableFormat.Borders.Color = Color.get_Black()
        row = table.AddRow(False, 1)
        row.Height = 20.0
        section.Tables.Add(table)

    def remove_signature(file_path):
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            para.text = para.text.replace("Evaluation Warning: The document was created with Spire.Doc for Python.", "")
            
        doc.save(file_path)
        
    output_path = "./worksheet_output"
    if not os.path.exists(output_path):
        os.makedirs(output_path)

    file_name = file_name.split(".")[0]
    questions = json.loads(questions)
    
    document = Document()
    section = document.AddSection()

    # generate header
    header = section.HeadersFooters.Header
    header_paragraph = header.AddParagraph()
    header_paragraph.AppendText(f'{file_name} Exercises')
    style = ParagraphStyle(document)
    style.CharacterFormat.FontSize = 26
    document.Styles.Add(style)
    header_paragraph.ApplyStyle(style.Name)

    for question in questions:
        # generate the question heading
        question_title = f'Exercise {question["number"]}: {question["title"]}'
        title_paragraph = section.AddParagraph()
        title_paragraph.AppendBreak(BreakType.LineBreak)
        title_text = title_paragraph.AppendText(question_title)
        title_text.CharacterFormat.Bold = True
        title_paragraph.AppendBreak(BreakType.LineBreak)
        section.Body.Paragraphs.Add(title_paragraph)

        # generate the question description
        question_description = question["description"]
        desc_paragraph = section.AddParagraph()
        desc_paragraph.AppendText(question_description)
        desc_paragraph.AppendBreak(BreakType.LineBreak)
        section.Body.Paragraphs.Add(desc_paragraph)
    
        # generate code input table
        code_input = section.AddParagraph()
        code_input.AppendText("Code")
        section.Body.Paragraphs.Add(code_input)
        generate_table(section, document)

        # generate output input table
        output_input = section.AddParagraph()
        output_input.AppendText("Output")
        section.Body.Paragraphs.Add(output_input)
        generate_table(section, document)

    file = f'{file_name} Exercises.docx'
    file_path = os.path.join(output_path, file)
    document.SaveToFile(file_path, FileFormat.Docx)
    remove_signature(file_path)


if __name__ == "__main__":
    generated_files = {}
    txt_file = select_file()
    file_name = os.path.basename(txt_file)
    questions = generate_questions(txt_file)
    generate_worksheet(file_name, questions)
    forms_questions_json = generate_json_data(txt_file)
    generate_form(forms_questions_json)